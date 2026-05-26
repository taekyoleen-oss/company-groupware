"""
그룹웨어 첫 사용자 가이드 - Word 문서 생성 스크립트
출력: output/그룹웨어_사용자_가이드.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# --- 색상 토큰 ---
COLOR_PRIMARY = RGBColor(0x0E, 0x76, 0x90)      # Teal
COLOR_DARK    = RGBColor(0x1F, 0x29, 0x37)      # Slate-800
COLOR_BODY    = RGBColor(0x33, 0x41, 0x55)      # Slate-700
COLOR_MUTED   = RGBColor(0x64, 0x74, 0x8B)      # Slate-500
COLOR_ACCENT  = RGBColor(0xCA, 0x8A, 0x04)      # Amber-600 (주의/팁)
COLOR_DANGER  = RGBColor(0xB9, 0x1C, 0x1C)      # Red-700
COLOR_OK      = RGBColor(0x15, 0x80, 0x3D)      # Green-700

DEFAULT_FONT = "맑은 고딕"


def set_run_font(run, font_name=DEFAULT_FONT, size=11, bold=False, color=COLOR_BODY):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_paragraph_border(paragraph, color_hex='0E7690', size=24, side='left'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr = OxmlElement(f'w:{side}')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), str(size))
    bdr.set(qn('w:space'), '8')
    bdr.set(qn('w:color'), color_hex)
    pBdr.append(bdr)
    pPr.append(pBdr)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=26, bold=True, color=COLOR_PRIMARY)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=False, color=COLOR_MUTED)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    add_paragraph_border(p, color_hex='0E7690', size=36, side='left')
    p.paragraph_format.left_indent = Cm(0.2)
    run = p.add_run(text)
    set_run_font(run, size=18, bold=True, color=COLOR_DARK)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, color=COLOR_PRIMARY)


def add_body(doc, text, size=11, bold=False, color=COLOR_BODY, indent_cm=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_step(doc, idx, text):
    """단계 표시: 1. 텍스트  (단계 번호 강조)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Cm(0.4)
    run_num = p.add_run(f"{idx}. ")
    set_run_font(run_num, size=11, bold=True, color=COLOR_PRIMARY)
    run = p.add_run(text)
    set_run_font(run, size=11, color=COLOR_BODY)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    mark = "•" if level == 0 else "–"
    run_m = p.add_run(f"{mark}  ")
    set_run_font(run_m, size=11, bold=True, color=COLOR_PRIMARY)
    run = p.add_run(text)
    set_run_font(run, size=11, color=COLOR_BODY)


def add_callout(doc, label, text, kind='tip'):
    """팁/주의/중요 박스 - 1셀 테이블 + 배경색"""
    palette = {
        'tip':     ('FFF7E6', 'CA8A04', '💡 팁'),
        'warn':    ('FEF2F2', 'B91C1C', '⚠️ 주의'),
        'info':    ('ECFEFF', '0E7690', 'ℹ️ 안내'),
        'success': ('ECFDF5', '15803D', '✅ 확인'),
    }
    bg, bar, default_label = palette.get(kind, palette['info'])
    final_label = label or default_label

    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, bg)

    # label
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(final_label)
    set_run_font(r1, size=11, bold=True, color=RGBColor.from_string(bar))

    # body
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.4
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5, color=COLOR_BODY)

    # 간격
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)


def add_kv_table(doc, rows, header=None):
    """2열 비교/요약 표"""
    tbl = doc.add_table(rows=(1 if header else 0) + len(rows), cols=2)
    tbl.autofit = True

    start = 0
    if header:
        hr = tbl.rows[0]
        for i, txt in enumerate(header):
            cell = hr.cells[i]
            shade_cell(cell, '0E7690')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            set_run_font(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        start = 1

    for i, (k, v) in enumerate(rows):
        row = tbl.rows[start + i]
        for j, txt in enumerate([k, v]):
            cell = row.cells[j]
            if j == 0:
                shade_cell(cell, 'F1F5F9')
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_run_font(r, size=10.5, bold=(j == 0), color=COLOR_DARK if j == 0 else COLOR_BODY)

    doc.add_paragraph()


def add_divider(doc):
    p = doc.add_paragraph()
    add_paragraph_border(p, color_hex='E2E8F0', size=6, side='bottom')


def build_document():
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # =============== 표지 ===============
    for _ in range(2):
        doc.add_paragraph()

    add_title(doc, "그룹웨어 사용자 가이드")
    add_subtitle(doc, "처음 사용하시는 분을 위한 단계별 안내서")

    # 메타 정보 박스
    meta = doc.add_table(rows=1, cols=1)
    cell = meta.cell(0, 0)
    shade_cell(cell, 'F8FAFC')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("본 문서는 그룹웨어의 핵심 기능 7가지를\n처음 사용하는 분도 따라 할 수 있도록 정리한 가이드입니다.")
    set_run_font(r, size=11, color=COLOR_MUTED)

    doc.add_paragraph()

    # 목차
    add_h2(doc, "📑 목차")
    toc_items = [
        "1. 시작하기 전에 — 알아두면 좋은 핵심 정보",
        "2. 로그인 및 비밀번호 변경",
        "3. 출근 기록하기 — 자동 인식 & 수동 확인",
        "4. 퇴근 시간 입력하기 — 야근/조퇴 처리",
        "5. 휴가 신청 및 취소",
        "6. 본인 정보(프로필) 확인",
        "7. 일정 게시 및 취소",
        "8. 자주 묻는 질문 (FAQ)",
    ]
    for item in toc_items:
        add_bullet(doc, item)

    doc.add_page_break()

    # =============== 1. 시작하기 전에 ===============
    add_h1(doc, "1. 시작하기 전에")
    add_body(doc, "본격적으로 앱을 사용하기 전에, 아래 정보를 먼저 확인해 주세요.")

    add_h2(doc, "✦ 사용자 역할(권한) 3단계")
    add_kv_table(doc, [
        ("실무자(Member)", "일반 직원. 본인 일정·휴가·출근만 관리합니다."),
        ("관리자(Manager)", "특정 직원의 결재자로 지정된 사용자. 본인 담당 직원의 휴가 결재 권한이 있습니다."),
        ("앱관리자(Super Admin)", "시스템 전체 관리자. 회원·팀·휴가일수·사무실 IP 등 모든 설정 권한을 갖습니다."),
    ], header=["구분", "설명"])

    add_callout(doc,
        "📌 본 회사의 앱관리자(웹관리자)",
        "본 그룹웨어의 앱관리자 계정은  manager@howdenkorea.com  입니다.\n"
        "휴가 결재가 막혀 있거나, 비밀번호 초기화 등 시스템 관련 문의는 이 계정으로 연락해 주세요.",
        kind='info')

    add_h2(doc, "✦ 어떤 화면에서 무엇을 할 수 있나요?")
    add_kv_table(doc, [
        ("캘린더", "팀/전사/개인 일정을 한 화면에서 확인하고 등록합니다."),
        ("공지사항", "회사 전체 또는 팀 단위 공지를 게시·열람합니다."),
        ("할 일(TO-DO)", "개인 작업을 체크리스트로 관리합니다."),
        ("프로필", "본인 정보·휴가 현황·출근 기록을 확인합니다."),
        ("결재함", "(관리자만) 담당 직원의 휴가 신청을 승인/반려합니다."),
    ], header=["메뉴", "기능"])

    doc.add_page_break()

    # =============== 2. 로그인 및 비밀번호 ===============
    add_h1(doc, "2. 로그인 및 비밀번호 변경")

    add_h2(doc, "STEP 1 — 최초 로그인")
    add_step(doc, 1, "회사에서 받은 이메일 주소와 임시 비밀번호로 로그인 페이지에 접속합니다.")
    add_step(doc, 2, "이메일과 비밀번호를 입력한 뒤 [로그인] 버튼을 누릅니다.")
    add_step(doc, 3, "최초 로그인 후 가입 승인 대기 상태일 경우, 웹관리자가 활성화할 때까지 잠시 기다려 주세요.")

    add_callout(doc, "⚠️ 가입 승인 대기 상태",
        "로그인 직후 '계정 승인 대기 중' 페이지가 보인다면, 앱관리자(manager@howdenkorea.com)에게 활성화를 요청해 주세요.",
        kind='warn')

    add_h2(doc, "STEP 2 — 비밀번호 변경")
    add_step(doc, 1, "화면 우측 상단의 본인 이름(또는 프로필 아이콘)을 클릭합니다.")
    add_step(doc, 2, "[프로필] 메뉴로 이동합니다.")
    add_step(doc, 3, "프로필 화면에서 [비밀번호 변경] 버튼을 누릅니다.")
    add_step(doc, 4, "현재 비밀번호 → 새 비밀번호 → 새 비밀번호 확인 순서로 입력합니다.")
    add_step(doc, 5, "[변경하기] 버튼을 누르면 즉시 적용됩니다.")

    add_callout(doc, "💡 비밀번호를 잊었을 때",
        "본인이 직접 초기화할 수 없습니다. 웹관리자(manager@howdenkorea.com)에게 문의하시면, 관리자 화면에서 "
        "비밀번호를  password  로 즉시 초기화해 드립니다. (이메일 발송 없음)\n"
        "  password 로 로그인한 직후, 반드시 프로필 화면에서 본인이 새 비밀번호로 변경해 주세요.",
        kind='tip')

    doc.add_page_break()

    # =============== 3. 출근 기록 ===============
    add_h1(doc, "3. 출근 기록하기")
    add_body(doc, "출근 기록은 두 가지 방식으로 처리됩니다. 어느 쪽이든 하루에 한 번만 기록되면 OK!")

    add_h2(doc, "방식 A — 자동 인식 (사무실 네트워크 사용 시)")
    add_step(doc, 1, "사무실 Wi-Fi/유선 네트워크에 연결된 상태로 그룹웨어에 접속합니다.")
    add_step(doc, 2, "로그인하면 시스템이 자동으로 사무실 IP를 인식합니다.")
    add_step(doc, 3, "화면 우측 상단에 '출근 완료' 표시(또는 출근 시각)가 자동으로 나타납니다.")

    add_callout(doc, "✅ 자동 인식이 잘 되었는지 확인하는 법",
        "헤더 우측 상단을 보세요. '출근 완료 09:02' 처럼 시간이 표시되어 있으면 정상 처리된 것입니다. "
        "별도의 버튼을 누를 필요가 없습니다.",
        kind='success')

    add_h2(doc, "방식 B — 수동 출근 확인 (재택/외근/네트워크 미인식 시)")
    add_step(doc, 1, "헤더 우측 상단의 [출근 확인] 버튼을 클릭합니다.")
    add_step(doc, 2, "팝업이 뜨면 [확인] 버튼을 눌러 출근을 기록합니다.")
    add_step(doc, 3, "기록 후 헤더 영역이 '출근 완료'로 바뀌면 정상 처리된 것입니다.")

    add_callout(doc, "ℹ️ 출근 시간은 자동으로 저장됩니다",
        "[출근 확인] 버튼을 누른 시각이 그대로 출근 시각으로 기록됩니다. 별도로 시각을 입력할 필요가 없습니다.",
        kind='info')

    add_callout(doc, "📌 앱관리자는 출근 기록 대상에서 제외됩니다",
        "앱관리자(manager@howdenkorea.com) 계정은 출근/휴가 통계에서 제외되어, 본인 화면에 [출근 확인] 버튼이 표시되지 않을 수 있습니다.",
        kind='info')

    doc.add_page_break()

    # =============== 4. 퇴근 입력 ===============
    add_h1(doc, "4. 퇴근 시간 입력하기")
    add_body(doc, "퇴근 시간은 따로 입력하지 않으면 기본값(18:00)으로 저장됩니다. "
                  "야근·조퇴 등 다른 시각으로 기록하고 싶을 때만 별도로 입력하세요.")

    add_h2(doc, "퇴근 시각을 직접 입력해야 하는 경우")
    add_bullet(doc, "18시 이후까지 야근한 경우")
    add_bullet(doc, "18시 이전에 조퇴한 경우")
    add_bullet(doc, "정확한 근무 시간 기록이 필요한 경우")

    add_h2(doc, "STEP — 퇴근 입력 방법")
    add_step(doc, 1, "헤더 우측 상단의 출근 표시 영역(또는 [퇴근 입력] 버튼)을 클릭합니다.")
    add_step(doc, 2, "퇴근 시각을 선택하거나 직접 입력합니다. (예: 20:30)")
    add_step(doc, 3, "[저장] 버튼을 눌러 기록합니다.")

    add_callout(doc, "⚠️ 미입력 시 자동 처리 규칙",
        "퇴근 시각을 입력하지 않고 하루가 끝나면, 시스템이 자동으로 그날의 퇴근 시각을 18:00으로 고정 저장합니다. "
        "야근한 날에는 반드시 직접 퇴근 시각을 기록해 주세요.",
        kind='warn')

    add_callout(doc, "💡 퇴근 후에도 수정 가능",
        "당일 내라면 이미 저장된 퇴근 시각도 다시 수정할 수 있습니다. "
        "이전 날짜의 기록 수정이 필요하면 웹관리자에게 요청해 주세요.",
        kind='tip')

    doc.add_page_break()

    # =============== 5. 휴가 신청/취소 ===============
    add_h1(doc, "5. 휴가 신청 및 취소")
    add_body(doc, "휴가는 본인의 결재자(관리자) 지정 여부에 따라 처리 흐름이 달라집니다.")

    add_h2(doc, "STEP 1 — 휴가 신청하기")
    add_step(doc, 1, "캘린더 화면에서 휴가를 사용할 날짜를 클릭합니다.")
    add_step(doc, 2, "일정 등록 모달에서 '휴가' 유형을 선택합니다.")
    add_step(doc, 3, "휴가 종류(연차/반차/오전반차/오후반차 등)와 사유를 입력합니다.")
    add_step(doc, 4, "[신청] 버튼을 누르면 결재 흐름이 자동으로 시작됩니다.")

    add_h2(doc, "✦ 신청 후 처리 흐름")
    add_kv_table(doc, [
        ("관리자가 지정된 경우",
         "지정된 관리자에게 결재 요청이 전달됩니다. 승인 시 캘린더에 정식 표시됩니다."),
        ("관리자가 지정되지 않은 경우",
         "본인이 본인 결재자인 경우 → 신청과 동시에 '자동 승인'으로 처리됩니다."),
        ("관리자 본인의 휴가 신청",
         "관리자라 하더라도 본인의 휴가는 본인이 결재할 수 없으며, 앱관리자(manager@howdenkorea.com)가 처리합니다."),
    ], header=["상황", "처리 방식"])

    add_callout(doc, "✅ 신청 직후 캘린더에 휴가가 보인다면?",
        "본인이 본인 결재자(자동 승인 대상)인 경우입니다. 별도 승인 없이 즉시 확정된 것입니다.",
        kind='success')

    add_h2(doc, "STEP 2 — 휴가 취소(삭제) 하기")
    add_step(doc, 1, "캘린더에서 본인이 신청한 휴가 일정을 클릭합니다.")
    add_step(doc, 2, "휴가 상세 화면에서 [취소 신청] 또는 [삭제] 버튼을 누릅니다.")
    add_step(doc, 3, "취소 사유를 입력하고 [신청] 버튼을 누릅니다.")

    add_callout(doc, "⚠️ 휴가 취소는 '자동 승인'되지 않습니다",
        "이미 승인된 휴가의 취소는 반드시 관리자(또는 앱관리자)의 별도 승인이 필요합니다.\n"
        "  • 본인의 결재자가 지정되어 있다면 → 해당 관리자가 승인\n"
        "  • 결재자가 없다면 → 앱관리자(manager@howdenkorea.com)가 승인\n"
        "취소 신청을 했더라도 승인 전까지 일정은 그대로 유지됩니다.",
        kind='warn')

    add_h2(doc, "✦ 잔여 휴가일수는 어디서 확인하나요?")
    add_body(doc, "→ 헤더 우측의 본인 이름 클릭 → [프로필] 메뉴에서 확인 가능합니다. (다음 섹션 참고)")

    doc.add_page_break()

    # =============== 6. 프로필 ===============
    add_h1(doc, "6. 본인 정보(프로필) 확인")

    add_h2(doc, "프로필 화면 진입 방법")
    add_step(doc, 1, "화면 우측 상단에 표시된 본인 이름(또는 프로필 원형 아이콘)을 클릭합니다.")
    add_step(doc, 2, "드롭다운 메뉴에서 [프로필] 항목을 선택합니다.")

    add_callout(doc, "💡 본인 이름이 안 보일 때",
        "모바일 화면에서는 우측 상단의 햄버거(≡) 메뉴 또는 하단 탭의 [프로필]을 통해 진입할 수 있습니다.",
        kind='tip')

    add_h2(doc, "프로필에서 확인할 수 있는 정보")
    add_bullet(doc, "기본 정보 — 이름, 이메일, 소속 팀, 직책")
    add_bullet(doc, "휴가 현황 — 총 휴가일수 / 사용 일수 / 잔여 일수")
    add_bullet(doc, "휴가 기록 — 연도별 휴가 신청·승인·취소 이력")
    add_bullet(doc, "출근 기록 — 월별 출근/퇴근 시각 이력")
    add_bullet(doc, "비밀번호 변경 — 현재 비밀번호로 즉시 변경 가능")

    doc.add_page_break()

    # =============== 7. 일정 게시 및 취소 ===============
    add_h1(doc, "7. 일정 게시 및 취소")

    add_h2(doc, "STEP 1 — 일정 등록하기")
    add_step(doc, 1, "[캘린더] 메뉴로 이동합니다.")
    add_step(doc, 2, "원하는 날짜의 빈 영역을 클릭하거나, 우측 상단의 [+ 일정 추가] 버튼을 누릅니다.")
    add_step(doc, 3, "일정 입력 모달에서 아래 항목을 채웁니다.")
    add_bullet(doc, "제목 — 일정명", level=1)
    add_bullet(doc, "유형 — 일반 일정 / 휴가 / 회의 등", level=1)
    add_bullet(doc, "공개 범위 — 전사(company) / 팀(team) / 개인(private)", level=1)
    add_bullet(doc, "시작·종료 일시", level=1)
    add_bullet(doc, "장소·메모(선택)", level=1)
    add_step(doc, 4, "[등록] 버튼을 누르면 캘린더에 즉시 게시됩니다.")

    add_callout(doc, "ℹ️ 공개 범위 차이",
        "• 전사(company) → 모든 직원이 볼 수 있음\n"
        "• 팀(team) → 같은 팀 구성원만 볼 수 있음\n"
        "• 개인(private) → 본인만 볼 수 있음 (관리자도 열람 불가)",
        kind='info')

    add_h2(doc, "STEP 2 — 일정 수정")
    add_step(doc, 1, "캘린더에서 수정할 일정을 클릭합니다.")
    add_step(doc, 2, "상세 화면에서 [수정] 버튼을 누르고 내용을 변경한 뒤 [저장]합니다.")

    add_h2(doc, "STEP 3 — 일정 취소(삭제)")
    add_step(doc, 1, "캘린더에서 삭제할 일정을 클릭합니다.")
    add_step(doc, 2, "상세 화면 하단의 [삭제] 버튼을 누릅니다.")
    add_step(doc, 3, "확인 팝업에서 [삭제] 버튼을 다시 눌러 확정합니다.")

    add_callout(doc, "⚠️ 일반 일정 vs 휴가 일정의 차이",
        "• 일반 일정 → 본인이 작성한 일정은 언제든 직접 삭제 가능합니다.\n"
        "• 휴가 일정 → '5. 휴가 신청 및 취소' 섹션 참고. 별도의 취소 결재가 필요합니다.",
        kind='warn')

    doc.add_page_break()

    # =============== 8. FAQ ===============
    add_h1(doc, "8. 자주 묻는 질문 (FAQ)")

    faqs = [
        ("Q. 비밀번호를 잊어버렸어요.",
         "A. 본인이 직접 재설정할 수는 없습니다. 앱관리자(manager@howdenkorea.com)에게 연락하시면 "
         "관리자 화면의 [비밀번호 초기화] 기능으로 비밀번호를 'password' 로 변경해 드립니다. "
         "이메일은 발송되지 않으니, 변경 완료 후 'password' 로 로그인한 뒤 프로필 화면에서 즉시 새 비밀번호로 변경하세요."),
        ("Q. 사무실에 있는데 자동 출근 처리가 안 돼요.",
         "A. 네트워크가 사무실 Wi-Fi에 연결되어 있는지 확인하세요. 그래도 안 되면 헤더 우측 상단의 [출근 확인] 버튼을 직접 눌러 수동 처리하시면 됩니다."),
        ("Q. 야근했는데 퇴근 입력을 깜빡했어요.",
         "A. 미입력 시 그날의 퇴근 시각은 18:00으로 자동 고정 저장됩니다. 정확한 기록이 필요하다면 앱관리자에게 수정 요청해 주세요."),
        ("Q. 휴가 신청을 했는데 결재자가 누구인지 모르겠어요.",
         "A. 프로필 화면 → 휴가 기록에서 본인 결재자 정보를 확인할 수 있습니다. 결재자가 지정되어 있지 않다면 자동 승인 또는 앱관리자 승인 대상입니다."),
        ("Q. 휴가 취소가 왜 바로 안 되나요?",
         "A. 이미 확정된 일정의 변경은 신중함이 필요해, 관리자(또는 앱관리자)의 별도 승인 절차를 거치도록 설계되어 있습니다."),
        ("Q. 본인 이름을 눌렀는데 프로필 메뉴가 안 보여요.",
         "A. 모바일 환경에서는 하단 탭 또는 햄버거 메뉴(≡)에서 [프로필]을 찾을 수 있습니다."),
        ("Q. 다른 사람의 휴가 일정도 보고 싶어요.",
         "A. 캘린더에서 공개 범위가 '전사' 또는 '같은 팀'으로 설정된 휴가는 확인할 수 있습니다. 개인(private) 일정은 작성자 본인만 볼 수 있습니다."),
        ("Q. 시스템 관련 문의는 어디로 하면 되나요?",
         "A. 본 그룹웨어의 앱관리자(웹관리자)는 manager@howdenkorea.com 입니다. 비밀번호 초기화·계정 활성화·휴가일수 조정 등 모든 시스템 요청을 이곳으로 보내주세요."),
    ]

    for q, a in faqs:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8)
        p_q.paragraph_format.space_after = Pt(2)
        r_q = p_q.add_run(q)
        set_run_font(r_q, size=11.5, bold=True, color=COLOR_DARK)

        p_a = doc.add_paragraph()
        p_a.paragraph_format.space_after = Pt(4)
        p_a.paragraph_format.line_spacing = 1.4
        p_a.paragraph_format.left_indent = Cm(0.4)
        r_a = p_a.add_run(a)
        set_run_font(r_a, size=11, color=COLOR_BODY)

    # =============== 마무리 ===============
    doc.add_paragraph()
    add_divider(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("📞  도움이 더 필요하신가요?")
    set_run_font(r, size=12, bold=True, color=COLOR_PRIMARY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("앱관리자 (웹관리자):  manager@howdenkorea.com")
    set_run_font(r2, size=11, color=COLOR_BODY)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("계정 활성화, 비밀번호 초기화, 휴가 결재, 출퇴근 기록 수정 등 모든 시스템 문의를 받습니다.")
    set_run_font(r3, size=10, color=COLOR_MUTED)

    return doc


def main():
    out_dir = os.path.join(os.path.dirname(__file__), "..", "output")
    os.makedirs(out_dir, exist_ok=True)
    base_name = "그룹웨어_사용자_가이드.docx"
    out_path = os.path.join(out_dir, base_name)

    doc = build_document()
    try:
        doc.save(out_path)
    except PermissionError:
        # 기존 파일이 Word 등에서 열려있을 때 → 타임스탬프 파일로 우회 저장
        from datetime import datetime
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        out_path = os.path.join(out_dir, f"그룹웨어_사용자_가이드_{ts}.docx")
        doc.save(out_path)

    abs_path = os.path.abspath(out_path)
    print(f"OK: {abs_path}")
    print(f"size: {os.path.getsize(abs_path):,} bytes")


if __name__ == "__main__":
    main()
